"""Pull readable text out of uploaded documents.

Without this, a PDF contributes only its filename to search: asking "what
languages are on my resume?" returned unrelated notes, because the resume's
contents had never been read. This is what turns "search your titles" into
"ask your documents".

Deliberately boring and dependency-light. No OCR, no layout reconstruction,
no table parsing — just the text layer, which is what an embedder wants
anyway. A scanned PDF has no text layer and is reported as such rather than
silently indexed as empty, because "we found nothing" and "there is nothing
to find" are different answers and the user deserves the honest one.
"""

import io
import logging
import re

log = logging.getLogger("vault.extract")

# Anything larger is refused rather than risked: extraction is CPU-bound and
# runs in the worker, so one pathological file would stall the queue.
MAX_EXTRACT_BYTES = 25 * 1024 * 1024

# Past this we stop reading. A 500-page book yields diminishing returns for
# retrieval and would dominate the index for its owner.
MAX_TEXT_CHARS = 400_000


class Unsupported(Exception):
    """The format has no text layer we can read — not an error, a fact."""


def _clean(text: str) -> str:
    # PDF extraction leaves ragged whitespace and hyphenated line breaks;
    # both hurt embedding quality more than they look like they should.
    text = text.replace("­", "")
    text = re.sub(r"-\n(\w)", r"\1", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        # An empty password unlocks many "protected" PDFs; a real one does not.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise Unsupported("This PDF is password-protected.") from exc

    out: list[str] = []
    for page in reader.pages:
        try:
            out.append(page.extract_text() or "")
        except Exception as exc:
            # One broken page must not lose the other four hundred.
            log.warning(f"Skipped an unreadable PDF page: {type(exc).__name__}: {exc}")
        if sum(len(p) for p in out) > MAX_TEXT_CHARS:
            break
    return "\n\n".join(out)


def _from_docx(data: bytes) -> str:
    import docx

    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    # Tables carry real content in CVs and reports, and are invisible if you
    # only walk paragraphs.
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


_PLAIN_SUFFIXES = (".txt", ".md", ".markdown", ".csv", ".json", ".log", ".rst", ".yml", ".yaml")


def extract_text(data: bytes, filename: str = "", content_type: str = "") -> str:
    """Return readable text, or raise Unsupported with a reason a user can act on."""
    if len(data) > MAX_EXTRACT_BYTES:
        raise Unsupported(f"File is larger than {MAX_EXTRACT_BYTES // 1024 // 1024} MB.")

    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ctype:
        text = _from_pdf(data)
    elif name.endswith(".docx") or "wordprocessingml" in ctype:
        text = _from_docx(data)
    elif name.endswith(_PLAIN_SUFFIXES) or ctype.startswith("text/"):
        text = _from_plain(data)
    elif name.endswith(".doc"):
        # Legacy binary Word. Converting it needs LibreOffice; say so plainly
        # instead of indexing mojibake.
        raise Unsupported("Legacy .doc isn't readable — re-save it as .docx or PDF.")
    else:
        raise Unsupported("No text could be read from this file type.")

    text = _clean(text)[:MAX_TEXT_CHARS]
    if not text.strip():
        raise Unsupported(
            "No text layer found — this looks like a scan or an image-only PDF. "
            "Running it through OCR first would make it searchable."
        )
    return text


def chunk_text(text: str, size: int = 1200, overlap: int = 200) -> list[str]:
    """Split into overlapping windows, preferring paragraph then sentence breaks.

    The overlap matters: a fact that straddles a boundary is otherwise
    findable from neither side, which is the classic way RAG appears to
    "not know" something that is plainly in the document.
    """
    text = text.strip()
    if len(text) <= size:
        return [text] if text else []

    # Cap the overlap at half the window. An overlap approaching the size
    # makes the step approach one character: 5,000 characters became 4,901
    # near-identical chunks in testing, which terminates but floods the
    # embedder and bloats the index for a single document.
    overlap = max(0, min(overlap, size // 2))

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            window = text[start:end]
            # Break on the last paragraph, then sentence, then space — never
            # mid-word, which produces tokens the model has never seen.
            for sep in ("\n\n", ". ", "\n", " "):
                cut = window.rfind(sep)
                if cut > size // 2:
                    end = start + cut + len(sep)
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return chunks
